"""
Convert literature_survey.tex to a clean Word document.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEX_FILE = "/home/gk-krishnan/Desktop/VIT Paper/literature_survey.tex"
OUT_FILE = "/home/gk-krishnan/Desktop/VIT Paper/Literature_Survey.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def clean(text):
    """Strip common LaTeX markup from a string."""
    # citations → [AuthorYear]
    text = re.sub(r'\\cite\{([^}]+)\}', lambda m: '[' + m.group(1).split(',')[0].strip() + ']', text)
    # bold/italic/texttt
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)
    # math
    text = re.sub(r'\$([^$]+)\$', lambda m: m.group(1)
                  .replace('\\geq','≥').replace('\\leq','≤')
                  .replace('\\approx','≈').replace('\\times','×')
                  .replace('\\rho','ρ').replace('\\Delta','Δ')
                  .replace('\\pm','±').replace('\\alpha','α')
                  .replace('\\beta','β').replace('\\sigma','σ')
                  .replace('^{2}','²').replace('_{}','')
                  .replace('\\text{ccc}','ccc').replace('\\',''), text)
    # accented chars
    text = text.replace("V\\'{a}\\v{s}a", "Váša")
    text = text.replace("\\'{a}", "á").replace("\\`{a}", "à")
    text = text.replace("\\v{s}", "š").replace("\\'{e}", "é")
    text = text.replace("\\'e", "é").replace("\\'{i}", "í")
    text = text.replace("\\'{o}", "ó").replace("\\'{u}", "ú")
    text = text.replace("\\\"u", "ü").replace("\\\"o", "ö")
    text = text.replace("\\~{n}", "ñ")
    # special chars
    text = text.replace("---", "—").replace("--", "–")
    text = text.replace("``", "\u201c").replace("''", "\u201d")
    text = text.replace("`", "\u2018").replace("'", "\u2019")
    text = text.replace("~", " ").replace("\\%", "%").replace("\\&", "&")
    text = text.replace("\\,", " ").replace("\\ ", " ")
    text = text.replace("{,}", ",").replace("{.}", ".")
    # remove remaining braces
    text = re.sub(r'\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\s*', '', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def add_run_with_style(para, text, bold=False, italic=False, size=11):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return run


def set_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return para


def add_table_of_works(doc, rows):
    """Add the comparison table."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Table I: Comparison of Key Related Works")
    run.bold = True
    run.font.size = Pt(11)

    headers = ["Study", "Field", "Method", "Dataset / N", "Task", "Key Result"]
    table = doc.add_table(rows=1 + len(rows), cols=6)
    table.style = 'Table Grid'

    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            if ri == len(rows) - 1:   # our row — bold
                cells[ci].paragraphs[0].runs[0].bold = True

    return table


# ── parse .tex ────────────────────────────────────────────────────────────────

def parse_tex(path):
    with open(path) as f:
        src = f.read()

    # strip comments
    src = re.sub(r'%.*', '', src)
    # collapse whitespace lines
    src = re.sub(r'\n{3,}', '\n\n', src)
    return src


def build_docx(tex_src):
    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── title block ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Accessible Brain Morphometry from Point-of-Care\nUltra-Low-Field MRI Using Physics-Constrained Deep Learning")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Literature Survey — Related Work Section\nIEEE Journal of Biomedical and Health Informatics (JBHI)")
    sr.italic = True
    sr.font.size = Pt(11)

    doc.add_paragraph()

    # ── section heading ──
    set_heading(doc, "Related Work", level=1)

    # ── intro paragraph ──
    intro_match = re.search(r'\\section\{Related Work\}(.*?)\\subsection', tex_src, re.DOTALL)
    if intro_match:
        intro = clean(intro_match.group(1))
        if intro:
            p = doc.add_paragraph(intro)
            p.paragraph_format.first_line_indent = Inches(0.3)
            for run in p.runs:
                run.font.size = Pt(11)

    # ── subsections ──
    subsec_pattern = re.compile(
        r'\\subsection\{([^}]+)\}(.*?)(?=\\subsection|\\begin\{thebibliography\}|\\begin\{table)',
        re.DOTALL
    )

    for m in subsec_pattern.finditer(tex_src):
        title_raw = clean(m.group(1))
        body      = m.group(2)

        set_heading(doc, title_raw, level=2)

        # subsubsections
        subsub_pattern = re.compile(
            r'\\subsubsection\{([^}]+)\}(.*?)(?=\\subsubsection|$)',
            re.DOTALL
        )
        subsub_matches = list(subsub_pattern.finditer(body))

        if subsub_matches:
            # text before first subsubsection
            pre = body[:subsub_matches[0].start()].strip()
            if pre:
                cleaned = clean(pre)
                if cleaned and len(cleaned) > 20:
                    p = doc.add_paragraph(cleaned)
                    p.paragraph_format.first_line_indent = Inches(0.3)
                    for run in p.runs:
                        run.font.size = Pt(11)

            for sm in subsub_matches:
                set_heading(doc, clean(sm.group(1)), level=3)
                content = sm.group(2).strip()
                # split into paragraphs
                for para_text in re.split(r'\n\n+', content):
                    cleaned = clean(para_text)
                    if cleaned and len(cleaned) > 20:
                        p = doc.add_paragraph(cleaned)
                        p.paragraph_format.first_line_indent = Inches(0.3)
                        for run in p.runs:
                            run.font.size = Pt(11)
        else:
            for para_text in re.split(r'\n\n+', body):
                # handle enumerate
                if '\\begin{enumerate}' in para_text:
                    items = re.findall(r'\\item\s+(.*?)(?=\\item|\\end\{enumerate\})', para_text, re.DOTALL)
                    for item in items:
                        cleaned = clean(item)
                        if cleaned:
                            p = doc.add_paragraph(cleaned, style='List Number')
                            for run in p.runs:
                                run.font.size = Pt(11)
                else:
                    cleaned = clean(para_text)
                    if cleaned and len(cleaned) > 20:
                        p = doc.add_paragraph(cleaned)
                        p.paragraph_format.first_line_indent = Inches(0.3)
                        for run in p.runs:
                            run.font.size = Pt(11)

    # ── comparison table ──
    table_rows = [
        ["Arnold et al. 2022",     "64 mT sim.",      "Gaussian blur + DenseNet-121",     "363 patients",           "Pathology detection",         "AUC 0.87–0.98"],
        ["Iglesias et al. 2023",   "64 mT real",      "SR-CNN + FreeSurfer",              "24 neurological pts",    "Brain morphometry",           "Hippocampus r=0.85; cerebrum r=0.92"],
        ["Dayarathna et al. 2024", "64 mT real",      "Adversarial diffusion + attention","45 subjects",            "ULF→HF translation",          "PSNR 24.04 dB, SSIM 0.907"],
        ["Váša et al. 2025",       "64 mT real",      "SynthSeg+ on MRR volume",          "23 adults",              "Morphometry reliability",     "ICC 0.93–0.97; Pearson r=0.88–0.94"],
        ["Islam et al. 2025",      "64 mT real",      "SynthSR / LoHiResGAN",             "92 healthy adults",      "Volumetric consistency",      "LoHiResGAN: ΔICV = +0.89%"],
        ["Hsu et al. 2025",        "64 mT real",      "SynthSeg+ on TomoBrain",           "60 adults",              "Morphometry protocol",        "TomoBrain best accuracy"],
        ["Baljer et al. 2025",     "64 mT real",      "Multi-orientation U-Net",          "56 infants",             "Pediatric super-resolution",  "Deep brain r=0.94, CCC=0.94"],
        ["Gopinath et al. 2025",   "LF simulated",    "3D U-Net (Recon-Any)",             "Paired HF/LF scans",     "Cortical surface mapping",    "Surface area r=0.96; thickness r=0.70"],
        ["Javadi et al. 2025",     "64 mT sim.",      "SR3 diffusion model",              "BraTS 2019 sim.",        "ULF→HF translation",          "SSIM>0.97; preserves pathology"],
        ["Ringshaw et al. 2026",   "64 mT real",      "MiniMORPH segmentation",           "78 infants",             "Infant brain volumes",        "ICV r=0.96, putamen r=0.97"],
        ["Rebsamen et al. 2020",   "3T",              "DL+DiReCT",                        "OASIS-3 (N=2,643)",      "Cortical thickness",          "r=0.887 vs. FreeSurfer"],
        ["Hasegawa et al. 2025",   "3T",              "3D ResNet-34 multitask",           "153 dementia patients",  "CSF Aβ prediction",           "RMSE 500 pg/mL (Aβ42)"],
        ["OUR WORK",               "64 mT sim.+real", "Physics-constrained ViT3D",        "IXI n=156; OASIS n=375; real n=23", "nWBV prediction", "r=0.892 [CI: 0.801–0.943]; age ρ=−0.597, p=0.003"],
    ]
    add_table_of_works(doc, table_rows)

    # ── references ──
    doc.add_paragraph()
    set_heading(doc, "References", level=2)

    bib_match = re.search(r'\\begin\{thebibliography\}\{[^}]+\}(.*?)\\end\{thebibliography\}', tex_src, re.DOTALL)
    if bib_match:
        bib_body = bib_match.group(1)
        items = re.split(r'\\bibitem\{[^}]+\}', bib_body)
        ref_num = 1
        for item in items:
            item = item.strip()
            if not item:
                continue
            # join lines, collapse whitespace
            item = re.sub(r'\s+', ' ', item)
            cleaned = clean(item)
            if cleaned and len(cleaned) > 10:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent       = Inches(0.4)
                p.paragraph_format.first_line_indent = Inches(-0.4)
                r = p.add_run(f"[{ref_num}] ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(cleaned)
                r2.font.size = Pt(10)
                ref_num += 1

    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")
    return OUT_FILE


if __name__ == "__main__":
    src = parse_tex(TEX_FILE)
    build_docx(src)

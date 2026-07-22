"""Generate a clean 300-word project summary Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path("/home/gk-krishnan/Desktop/VIT Paper/Project_Summary.docx")

doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.25)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Accessible Brain Morphometry from Point-of-Care\nUltra-Low-Field MRI Using Physics-Constrained Deep Learning")
r.bold = True; r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.add_paragraph()

def section(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x7A, 0x8A)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p

section(doc, "What We Built")
body(doc, (
    "A deep learning pipeline that predicts normalized whole brain volume (nWBV) — "
    "a clinical biomarker for dementia — directly from a 64 mT Hyperfine bedside MRI scanner. "
    "These scanners cost roughly 10× less than hospital 3T MRI systems and require no shielded room, "
    "making them suitable for low-resource clinics and point-of-care settings."
))

section(doc, "How It Works")
body(doc, (
    "Stage 1: A 3D Vision Transformer (ViT3D, 4.23M parameters) is pre-trained on physics-simulated "
    "64 mT data generated from 156 healthy subjects (IXI dataset). The simulation uses real measured "
    "MR physics at 64 mT — tissue-specific T1/T2 relaxation times, Rician noise, and B0 field "
    "inhomogeneity — rather than simple Gaussian blur used by prior work. "
    "Stage 2: The model is fine-tuned on 375 OASIS-1 subjects with FreeSurfer-derived nWBV labels "
    "and Clinical Dementia Rating (CDR) scores. All layers are updated at a low learning rate (5×10⁻⁵)."
))

section(doc, "Results")
body(doc, (
    "On the OASIS-1 test set (n=38), the model achieves Pearson r = 0.892 [95% CI: 0.801–0.943] "
    "for nWBV prediction — comparable to established cortical morphometry pipelines. The model "
    "correctly orders subjects by dementia severity without CDR as a training label: "
    "CDR 0.0 → 0.759, CDR 0.5 → 0.745, CDR 1.0 → 0.711 nWBV (−6.3% per CDR step, consistent "
    "with literature). On 23 real Hyperfine 64 mT scans (ds006557, Váša 2025), the model shows "
    "expected age-related brain atrophy (Spearman ρ = −0.597, p = 0.003). After lightweight domain "
    "adaptation using SynthSeg+ pseudo-labels, absolute error on real hardware is MAE = 0.010 nWBV "
    "units — within the clinical acceptability threshold of 0.02. Inference takes 130 ms per scan "
    "with no FreeSurfer dependency."
))

section(doc, "Novelty")
body(doc, (
    "This is the first work to apply a 3D Vision Transformer for direct nWBV regression from 64 mT MRI, "
    "the first to use physics-based 64 mT simulation (not Gaussian blur) for pre-training, and the "
    "first to validate on real Hyperfine hardware with FreeSurfer ground truth (via FastSurfer, n=23). "
    "Compared to SynthSeg+ (the leading segmentation-based method), our model achieves equivalent "
    "clinical utility at 130 ms vs 2–3 minutes, with no external software dependencies."
))

doc.save(OUT)
print(f"Saved: {OUT}")
